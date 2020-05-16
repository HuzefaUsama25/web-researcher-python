import bs4
import requests
import sys
import re 
import unicodedata
import os
import random
import datetime
import linecache
import time
import urllib3
from docx import Document

#create a for-loop which scrapes the links one by one from the links.txt file
for linenum in range (0,100):
    try:
        #open link num 1-49 from links.txt
        linenumthlink = str(linecache.getline('links.txt',linenum))
        if ('youtube' in linenumthlink):
            print("Skipping link which does not allow scraping....")
        else:
            #print the current number of website being scraped along with it's url
            print(str(str(linenum)+" "+str(linenumthlink)))
            #give a name to the file in which current scrapes are going to be put
            filename=r"D:\Huzefa\Desktop\The Big Researcher\\"+str(linenum)+".txt"
            #get request for current website from the links.txt file
            headers = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64;     x64; rv:66.0) Gecko/20100101 Firefox/66.0", "Accept-Encoding":"gzip, deflate",     "Accept":"text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8", "DNT":"1","Connection":"close", "Upgrade-Insecure-Requests":"1"}
            res = requests.get(linenumthlink, headers=headers)
            #convert the current website to BeautifulSoup object
            #res.raise_for_status()
            soup = bs4.BeautifulSoup(res.text, "lxml")
            title = soup.select("title")
            titletext = title[0].getText()
            #open the file in which the current scrapes are to be put
            filename=r"D:\Huzefa\Desktop\The Big Researcher\\"+str(linenum)+" - "+str(titletext)+".txt"
            if ("404" in titletext) or ("page not found" in titletext):
                print("PAGE NOT FOUND")
            else:
                linenum = Document()
                linenum.add_heading(titletext, 0)
                #loop through the websites html code and write each paragraph to the text file and then leave a line
                for i in soup.select("p"):
                   f = i.text
                   #f = (unicodedata.normalize('NFD', re.sub("[\(\[].*?[\)\]]", "", f)).encode('ascii', 'ignore'))
                   linenum.add_paragraph( f , style='Subtitle')
                   location = r"D:\Huzefa\Desktop\The Big Researcher\\"+titletext+".docx"          
                linenum.save(location)
                
                #sleep for 30 seconds so the unknown error does not occur
                #time.sleep(30)
    except:
        print("There was an error!")
        continue

